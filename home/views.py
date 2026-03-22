from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.template.loader import get_template
from django.conf import settings
import os


def landing(request):
    """Render the landing page."""
    if request.user.is_authenticated:
        return redirect('home')
    return render(request, 'landing.html')


@login_required
def home(request):
    """Render the home page with template selection."""
    template = request.GET.get('template', 'classic')
    return render(request, 'home.html', {'selected_template': template})


@login_required
def templates_gallery(request):
    """Render the templates gallery page showing all available resume templates."""
    return render(request, 'templates_gallery.html')


@login_required
def builder(request):
    """Render the resume builder form."""
    return render(request, 'index.html')


@login_required
def builder_fullstack(request):
    """Render the resume builder form tailored for Full Stack Developers."""
    context = {
        'profession': 'Full Stack Developer',
        'suggested_skills': [
            'JavaScript', 'Python', 'React', 'Node.js', 'Django', 'MongoDB',
            'PostgreSQL', 'Docker', 'AWS', 'Git', 'REST APIs', 'GraphQL'
        ],
        'suggested_projects': [
            'E-commerce Platform', 'Social Media Dashboard', 'Task Management App',
            'Real-time Chat Application', 'API Gateway Service'
        ],
        'template_type': 'modern',
        'form_title': 'Full Stack Developer CareerCraft'
    }
    return render(request, 'builder_fullstack.html', context)


@login_required
def builder_3d_generalist(request):
    """Render the resume builder form tailored for 3D Generalists."""
    context = {
        'profession': '3D Generalist',
        'suggested_skills': [
            'Maya', 'Blender', 'ZBrush', 'Substance Painter', 'Houdini',
            'Arnold', 'V-Ray', 'Unity', 'Unreal Engine', 'Marvelous Designer',
            'Rigging', 'Motion Capture'
        ],
        'suggested_projects': [
            'Character Modeling Portfolio', 'Environment Design', 'Product Visualization',
            'Animation Reel', 'VFX Simulation'
        ],
        'template_type': 'creative',
        'form_title': '3D Generalist CareerCraft'
    }
    return render(request, 'builder_3d.html', context)


@login_required
def builder_data_scientist(request):
    """Render the resume builder form tailored for Data Scientists."""
    context = {
        'profession': 'Data Scientist',
        'suggested_skills': [
            'Python', 'R', 'SQL', 'TensorFlow', 'PyTorch', 'Scikit-learn',
            'Pandas', 'NumPy', 'Tableau', 'Power BI', 'Spark', 'Hadoop',
            'Machine Learning', 'Deep Learning', 'Statistical Analysis'
        ],
        'suggested_projects': [
            'Predictive Analytics Model', 'NLP Sentiment Analysis', 'Recommendation System',
            'Computer Vision Project', 'Time Series Forecasting'
        ],
        'template_type': 'classic',
        'form_title': 'Data Scientist CareerCraft'
    }
    return render(request, 'builder_data.html', context)


@login_required
def builder_ux_designer(request):
    """Render the resume builder form tailored for UX/UI Designers."""
    context = {
        'profession': 'UX/UI Designer',
        'suggested_skills': [
            'Figma', 'Adobe XD', 'Sketch', 'InVision', 'Prototyping',
            'User Research', 'Wireframing', 'Design Systems', 'HTML/CSS',
            'Interaction Design', 'Usability Testing', 'Visual Design'
        ],
        'suggested_projects': [
            'Mobile App Redesign', 'Design System Creation', 'E-commerce UX Overhaul',
            'Dashboard Interface Design', 'User Research Study'
        ],
        'template_type': 'minimalist',
        'form_title': 'UX/UI Designer CareerCraft'
    }
    return render(request, 'builder_ux.html', context)


@login_required
def builder_executive(request):
    """Render the resume builder form tailored for Executives."""
    context = {
        'profession': 'Executive',
        'suggested_skills': [
            'Strategic Planning', 'Team Leadership', 'P&L Management',
            'Business Development', 'Stakeholder Management', 'Change Management',
            'M&A', 'Board Governance', 'Crisis Management', 'Investor Relations'
        ],
        'suggested_projects': [
            'Company Turnaround', 'IPO Launch', 'Merger Integration',
            'Digital Transformation', 'Market Expansion'
        ],
        'template_type': 'executive',
        'form_title': 'Executive CareerCraft'
    }
    return render(request, 'builder_executive.html', context)


@login_required
def gen_resume(request):
    """Generate and display the resume page."""
    if request.method == 'POST':
        template_type = request.POST.get('template_type', 'classic')

        context = {
            # Template type
            'template_type': template_type,
            # Profile
            'name': request.POST.get('name', ''),
            'title': request.POST.get('title', ''),
            'email': request.POST.get('email', ''),
            'phone': request.POST.get('phone', ''),
            'about': request.POST.get('about', ''),
            # Contact extras
            'location': request.POST.get('location', ''),
            'linkedin': request.POST.get('linkedin', ''),
            'github': request.POST.get('github', ''),
            'portfolio': request.POST.get('portfolio', ''),
            # Skills
            'prog_langs': request.POST.get('prog_langs', ''),
            'frameworks': request.POST.get('frameworks', ''),
            'databases': request.POST.get('databases', ''),
            'dev_tools': request.POST.get('dev_tools', ''),
            'web_tech': request.POST.get('web_tech', ''),
            # Experience 1
            'company1': request.POST.get('company1', ''),
            'post1': request.POST.get('post1', ''),
            'duration1': request.POST.get('duration1', ''),
            'location1': request.POST.get('location1', ''),
            'lin11': request.POST.get('lin11', ''),
            'achievements1': request.POST.get('achievements1', ''),
            # Experience 2
            'company2': request.POST.get('company2', ''),
            'post2': request.POST.get('post2', ''),
            'duration2': request.POST.get('duration2', ''),
            'location2': request.POST.get('location2', ''),
            'lin21': request.POST.get('lin21', ''),
            'achievements2': request.POST.get('achievements2', ''),
            # Experience 3
            'company3': request.POST.get('company3', ''),
            'post3': request.POST.get('post3', ''),
            'duration3': request.POST.get('duration3', ''),
            'location3': request.POST.get('location3', ''),
            'lin31': request.POST.get('lin31', ''),
            # Project 1
            'project1': request.POST.get('project1', ''),
            'durat1': request.POST.get('durat1', ''),
            'desc1': request.POST.get('desc1', ''),
            'project_tech1': request.POST.get('project_tech1', ''),
            'project_url1': request.POST.get('project_url1', ''),
            # Project 2
            'project2': request.POST.get('project2', ''),
            'durat2': request.POST.get('durat2', ''),
            'desc2': request.POST.get('desc2', ''),
            'project_tech2': request.POST.get('project_tech2', ''),
            'project_url2': request.POST.get('project_url2', ''),
            # Education 1
            'degree1': request.POST.get('degree1', ''),
            'college1': request.POST.get('college1', ''),
            'year1': request.POST.get('year1', ''),
            'gpa1': request.POST.get('gpa1', ''),
            'honors1': request.POST.get('honors1', ''),
            # Education 2
            'degree2': request.POST.get('degree2', ''),
            'college2': request.POST.get('college2', ''),
            'year2': request.POST.get('year2', ''),
            'gpa2': request.POST.get('gpa2', ''),
            'honors2': request.POST.get('honors2', ''),
            # Additional fields
            'languages': request.POST.get('languages', ''),
            'certifications': request.POST.get('certifications', ''),
            'interests': request.POST.get('interests', ''),
            # References
            'reference_name1': request.POST.get('reference_name1', ''),
            'reference_title1': request.POST.get('reference_title1', ''),
            'reference_contact1': request.POST.get('reference_contact1', ''),
            'reference_name2': request.POST.get('reference_name2', ''),
            'reference_title2': request.POST.get('reference_title2', ''),
            'reference_contact2': request.POST.get('reference_contact2', ''),
            # Board positions
            'board1': request.POST.get('board1', ''),
            'board_role1': request.POST.get('board_role1', ''),
            'board_year1': request.POST.get('board_year1', ''),
            'board2': request.POST.get('board2', ''),
            'board_role2': request.POST.get('board_role2', ''),
            'board_year2': request.POST.get('board_year2', ''),
            # Awards
            'award1': request.POST.get('award1', ''),
            'award_year1': request.POST.get('award_year1', ''),
            'award2': request.POST.get('award2', ''),
            'award_year2': request.POST.get('award_year2', ''),
        }

        # Select template based on type
        template_map = {
            'classic': 'resume.html',
            'modern': 'resume_modern.html',
            'minimalist': 'resume_minimalist.html',
            'creative': 'resume_creative.html',
            'executive': 'resume_executive.html',
        }

        template_name = template_map.get(template_type, 'resume.html')
        return render(request, template_name, context)
    return render(request, 'index.html')


@login_required
def export_resume_pdf(request):
    """Export resume as PDF."""
    if request.method == 'POST':
        # Get all form data
        context = get_resume_context(request.POST)
        template_type = request.POST.get('template_type', 'classic')
        
        # Render the resume template
        template_map = {
            'classic': 'resume.html',
            'modern': 'resume_modern.html',
            'minimalist': 'resume_minimalist.html',
            'creative': 'resume_creative.html',
            'executive': 'resume_executive.html',
        }
        template_name = template_map.get(template_type, 'resume.html')
        template = get_template(template_name)
        html = template.render(context)
        
        # Note: For production, you'll need to install weasyprint or xhtml2pdf
        # Example with weasyprint:
        # from weasyprint import HTML
        # pdf = HTML(string=html).write_pdf()
        # response = HttpResponse(pdf, content_type='application/pdf')
        # response['Content-Disposition'] = 'attachment; filename="resume.pdf"'
        # return response
        
        # For now, return HTML that can be printed as PDF
        response = HttpResponse(html, content_type='text/html')
        return response
    return redirect('builder')


@login_required
def export_resume_docx(request):
    """Export resume as DOCX/Word document."""
    if request.method == 'POST':
        # Get all form data
        context = get_resume_context(request.POST)
        name = context.get('name', 'resume').replace(' ', '_')
        
        # Create simple DOCX using python-docx
        try:
            from docx import Document
            doc = Document()
            
            # Add name and title
            doc.add_heading(context.get('name', ''), 0)
            if context.get('title'):
                doc.add_paragraph(context.get('title'), style='Subtitle')
            
            # Contact info
            contact_parts = []
            if context.get('email'):
                contact_parts.append(context.get('email'))
            if context.get('phone'):
                contact_parts.append(context.get('phone'))
            if context.get('location'):
                contact_parts.append(context.get('location'))
            if context.get('linkedin'):
                contact_parts.append(f"LinkedIn: {context.get('linkedin')}")
            if context.get('github'):
                contact_parts.append(f"GitHub: {context.get('github')}")
            if context.get('portfolio'):
                contact_parts.append(f"Portfolio: {context.get('portfolio')}")
            if contact_parts:
                doc.add_paragraph(' | '.join(contact_parts))
            
            # About
            if context.get('about'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(context.get('about'))
            
            # Experience
            for i in range(1, 4):
                company = context.get(f'company{i}', '')
                if company:
                    post = context.get(f'post{i}', '')
                    duration = context.get(f'duration{i}', '')
                    location_val = context.get(f'location{i}', '')
                    
                    heading = f'{post}' if post else ''
                    if company:
                        heading += f' at {company}' if heading else company
                    if heading:
                        doc.add_heading(heading, level=2)
                    if duration or location_val:
                        parts = []
                        if duration:
                            parts.append(duration)
                        if location_val:
                            parts.append(location_val)
                        doc.add_paragraph(' | '.join(parts))
                    
                    responsibilities = context.get(f'lin{i}1', '')
                    if responsibilities:
                        doc.add_paragraph(responsibilities)
                    
                    achievements = context.get(f'achievements{i}', '')
                    if achievements:
                        doc.add_paragraph('Key Achievements:')
                        for achievement in achievements.split('\n'):
                            if achievement.strip():
                                doc.add_paragraph(achievement.strip(), style='List Bullet')
            
            # Projects
            for i in range(1, 3):
                project = context.get(f'project{i}', '')
                if project:
                    desc = context.get(f'desc{i}', '')
                    tech = context.get(f'project_tech{i}', '')
                    url = context.get(f'project_url{i}', '')
                    durat = context.get(f'durat{i}', '')
                    
                    doc.add_heading(project, level=2)
                    if durat:
                        doc.add_paragraph(f'Duration: {durat}')
                    if desc:
                        doc.add_paragraph(desc)
                    if tech:
                        doc.add_paragraph(f'Technologies: {tech}')
                    if url:
                        doc.add_paragraph(f'URL: {url}')
            
            # Education
            for i in range(1, 3):
                degree = context.get(f'degree{i}', '')
                if degree:
                    college = context.get(f'college{i}', '')
                    year = context.get(f'year{i}', '')
                    gpa = context.get(f'gpa{i}', '')
                    honors = context.get(f'honors{i}', '')
                    
                    education_text = f'{degree}'
                    if college:
                        education_text += f' - {college}'
                    if year:
                        education_text += f' ({year})'
                    
                    doc.add_heading(education_text, level=2)
                    if gpa:
                        doc.add_paragraph(f'GPA: {gpa}')
                    if honors:
                        doc.add_paragraph(honors)
            
            # Skills
            skills_section = False
            skill_fields = [
                ('Programming Languages', 'prog_langs'),
                ('Frameworks & Libraries', 'frameworks'),
                ('Databases', 'databases'),
                ('Development Tools', 'dev_tools'),
                ('Web Technologies', 'web_tech'),
            ]
            
            for label, field in skill_fields:
                if context.get(field):
                    if not skills_section:
                        doc.add_heading('Skills', level=1)
                        skills_section = True
                    doc.add_paragraph(f'{label}: {context.get(field)}')
            
            # Additional sections
            if context.get('languages'):
                doc.add_heading('Languages', level=1)
                doc.add_paragraph(context.get('languages'))
            
            if context.get('certifications'):
                doc.add_heading('Certifications', level=1)
                for cert in context.get('certifications').split('\n'):
                    if cert.strip():
                        doc.add_paragraph(cert.strip(), style='List Bullet')
            
            if context.get('interests'):
                doc.add_heading('Interests', level=1)
                doc.add_paragraph(context.get('interests'))
            
            # References
            has_references = False
            for i in range(1, 3):
                ref_name = context.get(f'reference_name{i}', '')
                if ref_name:
                    if not has_references:
                        doc.add_heading('References', level=1)
                        has_references = True
                    ref_title = context.get(f'reference_title{i}', '')
                    ref_contact = context.get(f'reference_contact{i}', '')
                    ref_text = ref_name
                    if ref_title:
                        ref_text += f' - {ref_title}'
                    if ref_contact:
                        ref_text += f' ({ref_contact})'
                    doc.add_paragraph(ref_text)
            
            # Board Positions
            for i in range(1, 3):
                board = context.get(f'board{i}', '')
                if board:
                    role = context.get(f'board_role{i}', '')
                    year = context.get(f'board_year{i}', '')
                    
                    if not has_references:
                        doc.add_heading('Board Positions', level=1)
                        has_references = True  # Reuse flag
                    
                    board_text = f'{board}'
                    if role:
                        board_text += f' - {role}'
                    if year:
                        board_text += f' ({year})'
                    doc.add_paragraph(board_text)
            
            # Awards
            for i in range(1, 3):
                award = context.get(f'award{i}', '')
                if award:
                    year = context.get(f'award_year{i}', '')
                    
                    doc.add_heading('Awards & Honors', level=1)
                    award_text = award
                    if year:
                        award_text += f' ({year})'
                    doc.add_paragraph(award_text)
            
            # Save to response
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{name}_resume.docx"'
            return response
            
        except ImportError:
            messages.error(request, 'DOCX export requires python-docx package. Install with: pip install python-docx')
            return redirect('builder')
    
    return redirect('builder')


@login_required
def export_resume_txt(request):
    """Export resume as plain text."""
    if request.method == 'POST':
        context = get_resume_context(request.POST)
        name = context.get('name', 'resume').replace(' ', '_')
        
        # Build text content
        lines = []
        lines.append('=' * 80)
        lines.append(context.get('name', '').upper())
        lines.append('=' * 80)
        
        if context.get('title'):
            lines.append(context.get('title'))
        lines.append('')
        
        # Contact
        contact_parts = []
        if context.get('email'):
            contact_parts.append(f"Email: {context.get('email')}")
        if context.get('phone'):
            contact_parts.append(f"Phone: {context.get('phone')}")
        if context.get('location'):
            contact_parts.append(f"Location: {context.get('location')}")
        if context.get('linkedin'):
            contact_parts.append(f"LinkedIn: {context.get('linkedin')}")
        if context.get('github'):
            contact_parts.append(f"GitHub: {context.get('github')}")
        if contact_parts:
            lines.extend(contact_parts)
        lines.append('')
        
        # About
        if context.get('about'):
            lines.append('PROFESSIONAL SUMMARY')
            lines.append('-' * 80)
            lines.append(context.get('about'))
            lines.append('')
        
        # Experience
        for i in range(1, 4):
            company = context.get(f'company{i}', '')
            if company:
                post = context.get(f'post{i}', '')
                duration = context.get(f'duration{i}', '')
                location = context.get(f'location{i}', '')
                
                lines.append('')
                lines.append(f'EXPERIENCE: {post} at {company}')
                lines.append('-' * 80)
                lines.append(f'{duration} | {location}')
                lines.append('')
                
                if context.get(f'lin{i}1'):
                    lines.append('Responsibilities:')
                    lines.append(context.get(f'lin{i}1'))
                
                if context.get(f'achievements{i}'):
                    lines.append('')
                    lines.append('Key Achievements:')
                    for achievement in context.get(f'achievements{i}').split('\n'):
                        if achievement.strip():
                            lines.append(f'  • {achievement.strip()}')
        
        # Projects
        for i in range(1, 3):
            project = context.get(f'project{i}', '')
            if project:
                lines.append('')
                lines.append(f'PROJECT: {project}')
                lines.append('-' * 80)
                if context.get(f'durat{i}'):
                    lines.append(f'Duration: {context.get(f"durat{i}")}')
                if context.get(f'desc{i}'):
                    lines.append(context.get(f'desc{i}'))
                if context.get(f'project_tech{i}'):
                    lines.append(f'Technologies: {context.get(f"project_tech{i}")}')
                if context.get(f'project_url{i}'):
                    lines.append(f'URL: {context.get(f"project_url{i}")}')
        
        # Education
        for i in range(1, 3):
            degree = context.get(f'degree{i}', '')
            if degree:
                lines.append('')
                lines.append(f'EDUCATION: {degree}')
                lines.append('-' * 80)
                if context.get(f'college{i}'):
                    lines.append(f'{context.get(f"college{i}")}')
                if context.get(f'year{i}'):
                    lines.append(f'Year: {context.get(f"year{i}")}')
                if context.get(f'gpa{i}'):
                    lines.append(f'GPA: {context.get(f"gpa{i}")}')
                if context.get(f'honors{i}'):
                    lines.append(f'Honors: {context.get(f"honors{i}")}')
        
        # Skills
        skill_fields = [
            ('Programming Languages', 'prog_langs'),
            ('Frameworks', 'frameworks'),
            ('Databases', 'databases'),
            ('Development Tools', 'dev_tools'),
            ('Web Technologies', 'web_tech'),
        ]
        
        has_skills = any(context.get(field) for _, field in skill_fields)
        if has_skills:
            lines.append('')
            lines.append('SKILLS')
            lines.append('-' * 80)
            for label, field in skill_fields:
                if context.get(field):
                    lines.append(f'{label}: {context.get(field)}')
        
        # Additional sections
        if context.get('languages'):
            lines.append('')
            lines.append('LANGUAGES')
            lines.append('-' * 80)
            lines.append(context.get('languages'))
        
        if context.get('certifications'):
            lines.append('')
            lines.append('CERTIFICATIONS')
            lines.append('-' * 80)
            lines.append(context.get('certifications'))
        
        if context.get('interests'):
            lines.append('')
            lines.append('INTERESTS')
            lines.append('-' * 80)
            lines.append(context.get('interests'))
        
        lines.append('')
        lines.append('=' * 80)
        
        # Create response
        text_content = '\n'.join(lines)
        response = HttpResponse(text_content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{name}_resume.txt"'
        return response
    
    return redirect('builder')


def get_resume_context(post_data):
    """Helper function to extract resume context from POST data."""
    return {
        'name': post_data.get('name', ''),
        'title': post_data.get('title', ''),
        'email': post_data.get('email', ''),
        'phone': post_data.get('phone', ''),
        'about': post_data.get('about', ''),
        'location': post_data.get('location', ''),
        'linkedin': post_data.get('linkedin', ''),
        'github': post_data.get('github', ''),
        'portfolio': post_data.get('portfolio', ''),
        'prog_langs': post_data.get('prog_langs', ''),
        'frameworks': post_data.get('frameworks', ''),
        'databases': post_data.get('databases', ''),
        'dev_tools': post_data.get('dev_tools', ''),
        'web_tech': post_data.get('web_tech', ''),
        'company1': post_data.get('company1', ''),
        'post1': post_data.get('post1', ''),
        'duration1': post_data.get('duration1', ''),
        'location1': post_data.get('location1', ''),
        'lin11': post_data.get('lin11', ''),
        'achievements1': post_data.get('achievements1', ''),
        'company2': post_data.get('company2', ''),
        'post2': post_data.get('post2', ''),
        'duration2': post_data.get('duration2', ''),
        'location2': post_data.get('location2', ''),
        'lin21': post_data.get('lin21', ''),
        'achievements2': post_data.get('achievements2', ''),
        'company3': post_data.get('company3', ''),
        'post3': post_data.get('post3', ''),
        'duration3': post_data.get('duration3', ''),
        'location3': post_data.get('location3', ''),
        'lin31': post_data.get('lin31', ''),
        'project1': post_data.get('project1', ''),
        'durat1': post_data.get('durat1', ''),
        'desc1': post_data.get('desc1', ''),
        'project_tech1': post_data.get('project_tech1', ''),
        'project_url1': post_data.get('project_url1', ''),
        'project2': post_data.get('project2', ''),
        'durat2': post_data.get('durat2', ''),
        'desc2': post_data.get('desc2', ''),
        'project_tech2': post_data.get('project_tech2', ''),
        'project_url2': post_data.get('project_url2', ''),
        'degree1': post_data.get('degree1', ''),
        'college1': post_data.get('college1', ''),
        'year1': post_data.get('year1', ''),
        'gpa1': post_data.get('gpa1', ''),
        'honors1': post_data.get('honors1', ''),
        'degree2': post_data.get('degree2', ''),
        'college2': post_data.get('college2', ''),
        'year2': post_data.get('year2', ''),
        'gpa2': post_data.get('gpa2', ''),
        'honors2': post_data.get('honors2', ''),
        'languages': post_data.get('languages', ''),
        'certifications': post_data.get('certifications', ''),
        'interests': post_data.get('interests', ''),
        'reference_name1': post_data.get('reference_name1', ''),
        'reference_title1': post_data.get('reference_title1', ''),
        'reference_contact1': post_data.get('reference_contact1', ''),
        'reference_name2': post_data.get('reference_name2', ''),
        'reference_title2': post_data.get('reference_title2', ''),
        'reference_contact2': post_data.get('reference_contact2', ''),
        'board1': post_data.get('board1', ''),
        'board_role1': post_data.get('board_role1', ''),
        'board_year1': post_data.get('board_year1', ''),
        'board2': post_data.get('board2', ''),
        'board_role2': post_data.get('board_role2', ''),
        'board_year2': post_data.get('board_year2', ''),
        'award1': post_data.get('award1', ''),
        'award_year1': post_data.get('award_year1', ''),
        'award2': post_data.get('award2', ''),
        'award_year2': post_data.get('award_year2', ''),
    }


def login_view(request):
    """Handle user login."""
    if request.user.is_authenticated:
        return redirect('home')

    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        try:
            user = User.objects.get(email=email)
            user = authenticate(request, username=user.username, password=password)
            if user is not None:
                login(request, user)
                return redirect('home')
            else:
                messages.error(request, 'Invalid password.')
        except User.DoesNotExist:
            messages.error(request, 'No account found with this email.')

    return render(request, 'login.html')


def signup_view(request):
    """Handle user registration."""
    if request.user.is_authenticated:
        return redirect('home')

    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')

        if password != confirm_password:
            messages.error(request, 'Passwords do not match.')
            return render(request, 'signup.html')

        if User.objects.filter(email=email).exists():
            messages.error(request, 'An account with this email already exists.')
            return render(request, 'signup.html')

        username = email.split('@')[0]
        base_username = username
        counter = 1
        while User.objects.filter(username=username).exists():
            username = f"{base_username}{counter}"
            counter += 1

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            first_name=first_name,
            last_name=last_name
        )
        messages.success(request, 'Account created successfully! Please sign in.')
        return redirect('login')

    return render(request, 'signup.html')


def logout_view(request):
    """Handle user logout."""
    logout(request)
    return redirect('login')


@login_required
def profile(request):
    """Handle user profile dashboard."""
    # Get or create user profile for mobile number
    mobile = request.session.get('mobile', '')

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'update_profile':
            # Update personal information
            request.user.first_name = request.POST.get('first_name', '')
            request.user.last_name = request.POST.get('last_name', '')
            request.user.save()

            # Store mobile in session (since we're using default User model)
            mobile = request.POST.get('mobile', '')
            request.session['mobile'] = mobile

            messages.success(request, 'Profile updated successfully!')
            return redirect('profile')

        elif action == 'change_password':
            # Change password
            current_password = request.POST.get('current_password')
            new_password = request.POST.get('new_password')
            confirm_password = request.POST.get('confirm_password')

            if not request.user.check_password(current_password):
                messages.error(request, 'Current password is incorrect.')
                return redirect('profile')

            if new_password != confirm_password:
                messages.error(request, 'New passwords do not match.')
                return redirect('profile')

            request.user.set_password(new_password)
            request.user.save()

            # Re-authenticate user after password change
            user = authenticate(request, username=request.user.username, password=new_password)
            if user:
                login(request, user)

            messages.success(request, 'Password changed successfully!')
            return redirect('profile')

        elif action == 'delete_account':
            # Delete account
            password = request.POST.get('password')

            if not request.user.check_password(password):
                messages.error(request, 'Password is incorrect.')
                return redirect('profile')

            request.user.delete()
            messages.success(request, 'Your account has been deleted.')
            return redirect('landing')

    return render(request, 'profile.html', {'mobile': mobile})